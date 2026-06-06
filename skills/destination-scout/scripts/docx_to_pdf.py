#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_to_pdf.py — converte um levantamento .docx (gerado por md_to_docx.py) em PDF.

Uso:
    python3 docx_to_pdf.py entrada.docx "Saída.pdf"

Cadeia completa da skill (markdown -> docx -> pdf):
    python3 md_to_docx.py guia.md guia.docx
    python3 docx_to_pdf.py guia.docx "Guia.pdf"

Renderiza headings, subtítulo, parágrafos justificados, bullets e tabelas com
fonte DejaVu embarcada (acentos PT-BR ok). Os semáforos 🟢🟡🔴 viram bolinhas
coloridas (●) e emojis decorativos que a fonte não cobre são removidos.
Instala reportlab/python-docx se faltarem. Requer fontes DejaVu (padrão em Linux).
"""
import sys
import subprocess


def _ensure(mod, pkg=None):
    try:
        __import__(mod)
    except ImportError:
        print(f"Instalando {pkg or mod}...", file=sys.stderr)
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet", pkg or mod])


_ensure("docx", "python-docx")
_ensure("reportlab")

if len(sys.argv) < 3:
    print(__doc__)
    sys.exit(1)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import xml.etree.ElementTree as ET

FD = "/usr/share/fonts/truetype/dejavu"
pdfmetrics.registerFont(TTFont("DejaVu", f"{FD}/DejaVuSans.ttf"))
pdfmetrics.registerFont(TTFont("DejaVu-Bold", f"{FD}/DejaVuSans-Bold.ttf"))
pdfmetrics.registerFontFamily("DejaVu", normal="DejaVu", bold="DejaVu-Bold",
                              italic="DejaVu", boldItalic="DejaVu-Bold")

import re as _re
def emoji_fix(t):
    reps = {
        "🟢": '<font color="#1a7f37">●</font>',
        "🟡": '<font color="#d4a017">●</font>',
        "🔴": '<font color="#d11d1d">●</font>',
        "🌿": "·", "️": "",
    }
    for k, v in reps.items():
        t = t.replace(k, v)
    # remove emojis decorativos restantes que a DejaVu não cobre
    t = _re.sub(r'[\U0001F000-\U0001FAFF\U00002600-\U000027BF\U0001F1E6-\U0001F1FF]', '', t)
    return t

src, out = sys.argv[1], sys.argv[2]
doc = Document(src)
ACCENT = colors.HexColor("#2e5d3b")

ss = getSampleStyleSheet()
styles = {
    "Title": ParagraphStyle("T", parent=ss["Title"], fontName="DejaVu-Bold", textColor=ACCENT, fontSize=20, spaceBefore=6, spaceAfter=4, leading=24),
    "H1": ParagraphStyle("H1", parent=ss["Heading1"], fontName="DejaVu-Bold", textColor=ACCENT, fontSize=15, spaceBefore=22, spaceAfter=9, leading=18),
    "H2": ParagraphStyle("H2", parent=ss["Heading2"], fontName="DejaVu-Bold", textColor=colors.HexColor("#3d7a4e"), fontSize=12.5, spaceBefore=16, spaceAfter=6, leading=15),
    "H3": ParagraphStyle("H3", parent=ss["Heading3"], fontName="DejaVu-Bold", fontSize=11, spaceBefore=11, spaceAfter=4, leading=13),
    "Subtitle": ParagraphStyle("Sub", parent=ss["BodyText"], fontName="DejaVu", fontSize=9.5, textColor=colors.HexColor("#6b7280"), spaceAfter=10),
    "Body": ParagraphStyle("Body", parent=ss["BodyText"], fontName="DejaVu", fontSize=10, leading=14, alignment=TA_JUSTIFY, spaceAfter=6),
    "Bullet": ParagraphStyle("Bul", parent=ss["BodyText"], fontName="DejaVu", fontSize=10, leading=13, leftIndent=14, bulletIndent=4, spaceAfter=2),
    "Cell": ParagraphStyle("Cell", parent=ss["BodyText"], fontName="DejaVu", fontSize=8.5, leading=11),
    "CellH": ParagraphStyle("CellH", parent=ss["BodyText"], fontName="DejaVu-Bold", fontSize=8.5, leading=11, textColor=colors.white),
}

def esc(t):
    t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return emoji_fix(t)

def rich(para):
    """Converte runs (negrito/itálico) em markup reportlab."""
    out = ""
    for r in para.runs:
        t = esc(r.text)
        if r.bold:
            t = f"<b>{t}</b>"
        if r.italic:
            t = f"<i>{t}</i>"
        out += t
    return out or esc(para.text)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
flow = []

# itera blocos na ordem (parágrafos + tabelas) lendo o XML do body
body = doc.element.body
tbl_iter = iter(doc.tables)
par_iter = iter(doc.paragraphs)
tbls = list(doc.tables)
pars = list(doc.paragraphs)
ti = pi = 0
last_title = False

for child in body.iterchildren():
    tag = child.tag.split("}")[-1]
    if tag == "p":
        p = pars[pi]; pi += 1
        text = rich(p).strip()
        if not text:
            continue
        sty = p.style.name if p.style else ""
        if sty == "Title":
            flow.append(Paragraph(text, styles["Title"]))
            last_title = True
            continue
        elif sty.startswith("Heading 1") or sty == "Heading 1":
            flow.append(Paragraph(text, styles["H1"]))
        elif sty.startswith("Heading 2"):
            flow.append(Paragraph(text, styles["H2"]))
        elif sty.startswith("Heading"):
            flow.append(Paragraph(text, styles["H3"]))
        elif sty == "List Bullet":
            flow.append(Paragraph(text, styles["Bullet"], bulletText="•"))
        elif last_title:
            flow.append(Paragraph(text, styles["Subtitle"]))
        else:
            flow.append(Paragraph(text, styles["Body"]))
        last_title = False
    elif tag == "tbl":
        t = tbls[ti]; ti += 1
        data = []
        for ri, row in enumerate(t.rows):
            cells = []
            for cell in row.cells:
                txt = esc(cell.text.strip())
                cells.append(Paragraph(txt, styles["CellH"] if ri == 0 else styles["Cell"]))
            data.append(cells)
        ncols = len(data[0])
        avail = A4[0] - 3 * cm
        tbl = Table(data, colWidths=[avail / ncols] * ncols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), ACCENT),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b8ccbe")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f6f2")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        flow.append(Spacer(1, 4))
        flow.append(tbl)
        flow.append(Spacer(1, 8))

SimpleDocTemplate(out, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm,
                  leftMargin=1.5*cm, rightMargin=1.5*cm,
                  title=doc.core_properties.title or "").build(flow)
print(out)
