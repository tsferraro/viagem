#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_docx.py — converte um levantamento em Markdown para Word (.docx).

Uso:
    python3 md_to_docx.py entrada.md "Saída.docx"
    python3 md_to_docx.py entrada.md            # gera entrada.docx

Suporta: headings (# ## ###), tabelas pipe (| a | b |), bullets (- / *),
prosa em parágrafos justificados, negrito **assim** e itálico *assim*.
Instala python-docx se faltar. Pensado pra saída da skill destination-scout,
mas serve qualquer markdown simples.
"""
import sys
import re
import subprocess


def ensure_docx():
    try:
        import docx  # noqa
    except ImportError:
        print("Instalando python-docx...", file=sys.stderr)
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"])


def add_runs(paragraph, text):
    """Aplica **negrito** e *itálico* inline."""
    # divide preservando os delimitadores
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def is_table_row(line):
    return line.strip().startswith('|') and line.strip().endswith('|')


def is_table_sep(line):
    return bool(re.match(r'^\s*\|[\s:|-]+\|\s*$', line))


def parse_row(line):
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    ensure_docx()
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else re.sub(r'\.md$', '', src) + '.docx'

    with open(src, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    ACCENT = RGBColor(0x2e, 0x5d, 0x3b)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # vazio
        if not stripped:
            i += 1
            continue

        # tabela
        if is_table_row(line) and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            headers = parse_row(line)
            i += 2  # pula header + separador
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(parse_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = 'Light Grid Accent 1'
            for c, h in enumerate(headers):
                add_runs(t.rows[0].cells[c].paragraphs[0], h)
                for r in t.rows[0].cells[c].paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c, val in enumerate(row[:len(headers)]):
                    cells[c].paragraphs[0].text = ''
                    add_runs(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading('', level=0)
                add_runs(h, text)
            else:
                h = doc.add_heading('', level=min(level - 1, 4))
                add_runs(h, text)
            i += 1
            continue

        # bullets
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, text)
            i += 1
            continue

        # horizontal rule -> ignora
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = ACCENT
            i += 1
            continue

        # parágrafo de prosa (justificado)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        add_runs(p, stripped)
        i += 1

    doc.save(out)
    print(out)


if __name__ == '__main__':
    main()
